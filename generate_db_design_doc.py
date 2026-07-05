"""
MedAssist AI — Database Design Document Generator
Generates a professional Word (.docx) document with embedded ER diagrams.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

# ──────────────────────────────────────────────────────────────
# Output paths
# ──────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, "MedAssist_AI_Database_Design.docx")
DIAGRAMS_DIR = os.path.join(BASE_DIR, "diagrams")
os.makedirs(DIAGRAMS_DIR, exist_ok=True)


# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B4F72"):
    """Add a formatted table with coloured header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    return h


# ──────────────────────────────────────────────────────────────
# ER Diagram Generation (matplotlib)
# ──────────────────────────────────────────────────────────────
def draw_entity_box(ax, x, y, name, attrs, color="#2980B9", width=2.8, attr_h=0.22):
    """Draw a single entity box with attributes."""
    header_h = 0.35
    total_h = header_h + len(attrs) * attr_h + 0.1

    # Entity header
    rect = FancyBboxPatch((x, y), width, header_h,
                          boxstyle="round,pad=0.05",
                          facecolor=color, edgecolor="#1A5276", linewidth=1.5)
    ax.add_patch(rect)
    ax.text(x + width / 2, y + header_h / 2, name,
            ha='center', va='center', fontsize=8, fontweight='bold', color='white')

    # Attribute box
    attr_rect = FancyBboxPatch((x, y - len(attrs) * attr_h - 0.1), width,
                               len(attrs) * attr_h + 0.1,
                               boxstyle="round,pad=0.05",
                               facecolor="#EBF5FB", edgecolor="#1A5276", linewidth=1)
    ax.add_patch(attr_rect)

    for i, attr in enumerate(attrs):
        ay = y - (i + 1) * attr_h + 0.02
        prefix = "🔑 " if i == 0 else "    "
        ax.text(x + 0.1, ay, f"{prefix}{attr}", fontsize=6, va='center', color="#1A5276")

    return (x, y, x + width, y - len(attrs) * attr_h - 0.1)


def draw_relationship_line(ax, box1, box2, label, card1="1", card2="M"):
    """Draw a relationship line between two entity boxes."""
    x1 = (box1[0] + box1[2]) / 2
    y1 = (box1[1] + box1[3]) / 2
    x2 = (box2[0] + box2[2]) / 2
    y2 = (box2[1] + box2[3]) / 2

    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color="#2C3E50", lw=1.2))

    mid_x = (x1 + x2) / 2
    mid_y = (y1 + y2) / 2
    ax.text(mid_x, mid_y + 0.15, label, fontsize=5.5,
            ha='center', va='center', color="#8E44AD", fontstyle='italic',
            bbox=dict(boxstyle='round,pad=0.15', facecolor='#F5EEF8', edgecolor='#8E44AD', alpha=0.8))
    ax.text(x1 + (x2 - x1) * 0.15, y1 + (y2 - y1) * 0.15 + 0.1, card1,
            fontsize=6, ha='center', color='#E74C3C', fontweight='bold')
    ax.text(x1 + (x2 - x1) * 0.85, y1 + (y2 - y1) * 0.85 + 0.1, card2,
            fontsize=6, ha='center', color='#E74C3C', fontweight='bold')


def generate_module_diagram(title, entities, relationships, filename, figsize=(14, 10)):
    """Generate a single module ER diagram."""
    fig, ax = plt.subplots(1, 1, figsize=figsize)
    ax.set_xlim(-1, figsize[0] - 1)
    ax.set_ylim(-figsize[1] + 1, 2)
    ax.axis('off')
    ax.set_title(title, fontsize=14, fontweight='bold', color='#1B4F72', pad=20)

    boxes = {}
    for ename, (x, y, attrs, color) in entities.items():
        boxes[ename] = draw_entity_box(ax, x, y, ename, attrs, color)

    for (e1, e2, label, c1, c2) in relationships:
        if e1 in boxes and e2 in boxes:
            draw_relationship_line(ax, boxes[e1], boxes[e2], label, c1, c2)

    # Legend
    legend_items = [
        mpatches.Patch(color='#2980B9', label='Core Entity'),
        mpatches.Patch(color='#27AE60', label='Supporting Entity'),
        mpatches.Patch(color='#E67E22', label='Junction/Mapping Table'),
        mpatches.Patch(color='#8E44AD', label='Log/Audit Entity'),
    ]
    ax.legend(handles=legend_items, loc='lower right', fontsize=7, framealpha=0.9)

    plt.tight_layout()
    path = os.path.join(DIAGRAMS_DIR, filename)
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f"  ✅ Saved: {filename}")
    return path


def generate_all_diagrams():
    """Generate all ER diagrams."""
    diagrams = {}

    # ── 1. User Management ER Diagram ──
    diagrams["user_management_er"] = generate_module_diagram(
        "Module 1: User Management — Entity Relationship Diagram",
        {
            "User":           (1, 1, ["user_id (PK)", "email", "display_name", "role (FK)", "phone", "language_preference", "is_active", "created_at"], "#2980B9"),
            "UserRole":       (6, 1, ["role_id (PK)", "role_name", "role_description", "access_level", "is_active"], "#27AE60"),
            "UserPreference": (1, -2.5, ["preference_id (PK)", "user_id (FK)", "font_size", "high_contrast_mode", "theme", "timezone"], "#27AE60"),
            "SessionLog":     (6, -2.5, ["session_id (PK)", "user_id (FK)", "login_timestamp", "logout_timestamp", "device_info", "ip_address"], "#8E44AD"),
            "AccountStatus":  (10, 1, ["status_id (PK)", "status_name", "status_description"], "#27AE60"),
        },
        [
            ("UserRole", "User", "assigns role", "1", "M"),
            ("User", "UserPreference", "has preference", "1", "1"),
            ("User", "SessionLog", "logs sessions", "1", "M"),
            ("AccountStatus", "User", "applies status", "1", "M"),
        ],
        "01_user_management_er.png", figsize=(14, 7)
    )

    # ── 2. Patient Health Profile ER Diagram ──
    diagrams["patient_profile_er"] = generate_module_diagram(
        "Module 2: Patient Health Profile — Entity Relationship Diagram",
        {
            "User":             (0.5, 1, ["user_id (PK)", "email", "display_name", "role"], "#2980B9"),
            "PatientProfile":   (5, 1, ["patient_id (PK/FK)", "age", "gender", "blood_group", "condition", "allergies[]", "chronic_conditions[]", "assigned_doctor_id (FK)", "assigned_caregiver_id (FK)", "adherence", "status"], "#2980B9"),
            "EmergencyContact": (10, 1, ["contact_id (PK)", "patient_id (FK)", "contact_name", "contact_phone", "relationship", "is_primary"], "#27AE60"),
            "BloodGroup":       (5, -3.5, ["blood_group_id (PK)", "blood_group_name", "description"], "#27AE60"),
            "PatientStatus":    (10, -3.5, ["status_id (PK)", "status_name", "severity_level"], "#27AE60"),
        },
        [
            ("User", "PatientProfile", "has profile", "1", "1"),
            ("PatientProfile", "EmergencyContact", "has contacts", "1", "M"),
            ("BloodGroup", "PatientProfile", "classifies", "1", "M"),
            ("PatientStatus", "PatientProfile", "assigns status", "1", "M"),
        ],
        "02_patient_profile_er.png", figsize=(14, 8)
    )

    # ── 3. Doctor Professional ER Diagram ──
    diagrams["doctor_er"] = generate_module_diagram(
        "Module 3: Doctor Professional — Entity Relationship Diagram",
        {
            "User":           (0.5, 1, ["user_id (PK)", "email", "display_name", "role"], "#2980B9"),
            "DoctorProfile":  (5, 1, ["doctor_id (PK/FK)", "name", "specialization (FK)", "hospital", "experience_years", "qualification", "rating", "is_accepting_patients"], "#2980B9"),
            "ClinicalNote":   (10, 1, ["note_id (PK)", "doctor_id (FK)", "patient_id (FK)", "appointment_id (FK)", "note_content", "diagnosis_codes[]", "prescriptions[]"], "#27AE60"),
            "Specialization": (5, -2.5, ["specialization_id (PK)", "specialization_name", "specialization_description"], "#27AE60"),
            "AppointmentSlot":(10, -2.5, ["slot_id (PK)", "doctor_id (FK)", "date", "start_time", "end_time", "is_booked"], "#27AE60"),
        },
        [
            ("User", "DoctorProfile", "has profile", "1", "1"),
            ("Specialization", "DoctorProfile", "categorises", "1", "M"),
            ("DoctorProfile", "ClinicalNote", "writes notes", "1", "M"),
            ("DoctorProfile", "AppointmentSlot", "has slots", "1", "M"),
        ],
        "03_doctor_professional_er.png", figsize=(14, 7)
    )

    # ── 4. Caregiver Monitoring ER Diagram ──
    diagrams["caregiver_er"] = generate_module_diagram(
        "Module 4: Caregiver Monitoring — Entity Relationship Diagram",
        {
            "User":            (0.5, 1, ["user_id (PK)", "email", "display_name", "role"], "#2980B9"),
            "CaregiverProfile":(5, 1, ["caregiver_id (PK/FK)", "phone", "relationship_to_patient", "alert_preference_channel", "is_available"], "#2980B9"),
            "CaregiverPatientMapping": (10, 1, ["mapping_id (PK)", "caregiver_id (FK)", "patient_id (FK)", "relationship", "is_active", "alert_enabled"], "#E67E22"),
            "PatientProfile":  (10, -2, ["patient_id (PK)", "name", "condition", "status"], "#2980B9"),
            "RelationshipType":(5, -2, ["type_id (PK)", "type_name", "type_description"], "#27AE60"),
        },
        [
            ("User", "CaregiverProfile", "has profile", "1", "1"),
            ("CaregiverProfile", "CaregiverPatientMapping", "monitors via", "1", "M"),
            ("CaregiverPatientMapping", "PatientProfile", "maps to", "M", "1"),
            ("RelationshipType", "CaregiverPatientMapping", "classifies", "1", "M"),
        ],
        "04_caregiver_monitoring_er.png", figsize=(14, 7)
    )

    # ── 5. Medicine Management ER Diagram ──
    diagrams["medicine_er"] = generate_module_diagram(
        "Module 5: Medicine Management — Entity Relationship Diagram",
        {
            "PatientProfile":  (0.5, 1, ["patient_id (PK)", "name", "condition"], "#2980B9"),
            "MedicineSchedule":(5, 1, ["medicine_id (PK)", "patient_id (FK)", "name", "dosage", "frequency", "schedule_slots[]", "reminder_times[]", "is_active"], "#2980B9"),
            "MedicineLog":     (10, 1, ["log_id (PK)", "medicine_id (FK)", "patient_id (FK)", "status [taken/missed]", "slot", "timestamp"], "#8E44AD"),
            "WHOEssentialMedicine": (5, -2.5, ["eml_id (PK)", "medicine_name", "generic_name", "atc_code", "category", "formulation"], "#27AE60"),
            "MedicineCategory":(10, -2.5, ["category_id (PK)", "category_name", "parent_category_id"], "#27AE60"),
        },
        [
            ("PatientProfile", "MedicineSchedule", "has medicines", "1", "M"),
            ("MedicineSchedule", "MedicineLog", "generates logs", "1", "M"),
            ("MedicineCategory", "WHOEssentialMedicine", "categorises", "1", "M"),
        ],
        "05_medicine_management_er.png", figsize=(14, 7)
    )

    # ── 6. Appointment Booking ER Diagram ──
    diagrams["appointment_er"] = generate_module_diagram(
        "Module 6: Appointment Booking — Entity Relationship Diagram",
        {
            "PatientProfile": (0.5, 1, ["patient_id (PK)", "name", "phone"], "#2980B9"),
            "Appointment":    (5, 1, ["appointment_id (PK)", "patient_id (FK)", "doctor_id (FK)", "date", "time", "reason", "status", "notes"], "#2980B9"),
            "DoctorProfile":  (10, 1, ["doctor_id (PK)", "name", "specialization", "hospital"], "#2980B9"),
            "AppointmentStatus": (5, -2.5, ["status_id (PK)", "status_name", "is_terminal"], "#27AE60"),
            "AppointmentPolicy": (10, -2.5, ["policy_id (PK)", "policy_name", "advance_booking_days", "cancellation_window_hours"], "#27AE60"),
        },
        [
            ("PatientProfile", "Appointment", "books", "1", "M"),
            ("Appointment", "DoctorProfile", "assigned to", "M", "1"),
            ("AppointmentStatus", "Appointment", "classifies", "1", "M"),
            ("AppointmentPolicy", "Appointment", "governs", "1", "M"),
        ],
        "06_appointment_booking_er.png", figsize=(14, 7)
    )

    # ── 7. Emergency SOS ER Diagram ──
    diagrams["sos_er"] = generate_module_diagram(
        "Module 7: Emergency SOS — Entity Relationship Diagram",
        {
            "PatientProfile": (0.5, 1, ["patient_id (PK)", "name", "phone"], "#2980B9"),
            "SOSEvent":       (5.5, 1, ["sos_id (PK)", "patient_id (FK)", "timestamp", "location_lat", "location_lng", "status", "severity", "notified_contacts[]", "resolved_by"], "#2980B9"),
            "SOSStatus":      (5.5, -2, ["status_id (PK)", "status_name", "status_description"], "#27AE60"),
            "EmergencyService":(10.5, 1, ["service_id (PK)", "service_name", "service_phone", "service_type", "region"], "#27AE60"),
        },
        [
            ("PatientProfile", "SOSEvent", "triggers", "1", "M"),
            ("SOSStatus", "SOSEvent", "classifies", "1", "M"),
        ],
        "07_emergency_sos_er.png", figsize=(14, 6)
    )

    # ── 8. AI Chat and Voice ER Diagram ──
    diagrams["chat_er"] = generate_module_diagram(
        "Module 8: AI Chat and Voice — Entity Relationship Diagram",
        {
            "User":          (0.5, 1, ["user_id (PK)", "email", "display_name"], "#2980B9"),
            "ChatSession":   (5, 1, ["chat_id (PK)", "user_id (FK)", "language", "session_start", "message_count", "model_used"], "#2980B9"),
            "ChatMessage":   (10, 1, ["message_id (PK)", "chat_id (FK)", "role", "text", "sources[]", "emergency_flag", "timestamp"], "#27AE60"),
            "VoiceSession":  (5, -2.5, ["voice_id (PK)", "user_id (FK)", "language", "transcription", "ai_response", "model_used"], "#2980B9"),
            "RAGQueryLog":   (10, -2.5, ["query_log_id (PK)", "chat_id (FK)", "query_text", "sources_queried[]", "cache_hit", "query_duration_ms"], "#8E44AD"),
        },
        [
            ("User", "ChatSession", "initiates", "1", "M"),
            ("ChatSession", "ChatMessage", "contains", "1", "M"),
            ("User", "VoiceSession", "uses voice", "1", "M"),
            ("ChatMessage", "RAGQueryLog", "logs RAG query", "1", "1"),
        ],
        "08_ai_chat_voice_er.png", figsize=(14, 7)
    )

    # ── 9. Notification Engine ER Diagram ──
    diagrams["notification_er"] = generate_module_diagram(
        "Module 9: Notification Engine — Entity Relationship Diagram",
        {
            "User":          (0.5, 1, ["user_id (PK)", "email", "display_name"], "#2980B9"),
            "Notification":  (5, 1, ["notification_id (PK)", "recipient_id (FK)", "channel", "type", "title", "message", "priority", "is_read"], "#2980B9"),
            "NotifChannel":  (10, 1, ["channel_id (PK)", "channel_name", "api_provider", "is_active", "rate_limit"], "#27AE60"),
            "NotifPreference": (5, -2, ["preference_id (PK)", "user_id (FK)", "medicine_channel", "appointment_channel", "emergency_channel"], "#27AE60"),
        },
        [
            ("User", "Notification", "receives", "1", "M"),
            ("NotifChannel", "Notification", "delivers via", "1", "M"),
            ("User", "NotifPreference", "sets preference", "1", "1"),
        ],
        "09_notification_engine_er.png", figsize=(14, 6)
    )

    # ── 10. Alert Management ER Diagram ──
    diagrams["alert_er"] = generate_module_diagram(
        "Module 10: Alert Management — Entity Relationship Diagram",
        {
            "PatientProfile": (0.5, 1, ["patient_id (PK)", "name", "status"], "#2980B9"),
            "Alert":          (5, 1, ["alert_id (PK)", "type (FK)", "severity", "patient_id (FK)", "symptom", "status", "assigned_to (FK)"], "#2980B9"),
            "AlertType":      (10, 1, ["type_id (PK)", "type_name", "default_severity", "auto_notify"], "#27AE60"),
            "AlertSeverity":  (5, -2.5, ["severity_id (PK)", "severity_name", "severity_level", "response_time_mins"], "#27AE60"),
        },
        [
            ("PatientProfile", "Alert", "generates", "1", "M"),
            ("AlertType", "Alert", "classifies", "1", "M"),
            ("AlertSeverity", "Alert", "prioritises", "1", "M"),
        ],
        "10_alert_management_er.png", figsize=(14, 6)
    )

    # ── 11. Reports ER Diagram ──
    diagrams["reports_er"] = generate_module_diagram(
        "Module 11: Reports and Analytics — Entity Relationship Diagram",
        {
            "PatientProfile": (0.5, 1, ["patient_id (PK)", "name", "adherence"], "#2980B9"),
            "HealthReport":   (5, 1, ["report_id (PK)", "patient_id (FK)", "type", "data (JSON)", "generated_by", "date_range"], "#2980B9"),
            "AdherenceReport":(10, 1, ["adherence_report_id (PK)", "patient_id (FK)", "total_scheduled", "total_taken", "total_missed", "adherence_%"], "#2980B9"),
            "ReportType":     (5, -2, ["type_id (PK)", "type_name", "available_for_roles[]"], "#27AE60"),
        },
        [
            ("PatientProfile", "HealthReport", "has reports", "1", "M"),
            ("PatientProfile", "AdherenceReport", "has adherence", "1", "M"),
            ("ReportType", "HealthReport", "classifies", "1", "M"),
        ],
        "11_reports_analytics_er.png", figsize=(14, 6)
    )

    # ── 12. Audit & Compliance ER Diagram ──
    diagrams["audit_er"] = generate_module_diagram(
        "Module 12: Audit and Compliance — Entity Relationship Diagram",
        {
            "User":           (0.5, 1, ["user_id (PK)", "email", "role"], "#2980B9"),
            "AuditLog":       (5, 1, ["audit_id (PK)", "record_type", "record_id", "actor_id (FK)", "action", "data_hash", "block_hash", "block_index"], "#8E44AD"),
            "ConsentRecord":  (10, 1, ["consent_id (PK)", "patient_id (FK)", "consent_type", "granted_to", "consent_hash", "is_active"], "#2980B9"),
            "DataAccessLog":  (5, -2.5, ["access_log_id (PK)", "accessor_id (FK)", "patient_id (FK)", "data_type", "access_type", "timestamp"], "#8E44AD"),
        },
        [
            ("User", "AuditLog", "generates", "1", "M"),
            ("User", "DataAccessLog", "accesses data", "1", "M"),
            ("PatientProfile" if "PatientProfile" in {} else "User", "ConsentRecord", "grants consent", "1", "M"),
        ],
        "12_audit_compliance_er.png", figsize=(14, 7)
    )

    # ── MASTER ER DIAGRAM (overview) ──
    fig, ax = plt.subplots(1, 1, figsize=(20, 14))
    ax.set_xlim(-1, 19)
    ax.set_ylim(-13, 2)
    ax.axis('off')
    ax.set_title("MedAssist AI — Complete Entity Relationship Overview", fontsize=16, fontweight='bold', color='#1B4F72', pad=20)

    # Module boxes (simplified)
    modules = [
        ("User", 1, 1, ["user_id (PK)", "email", "role"], "#2980B9"),
        ("PatientProfile", 5, 1, ["patient_id (PK/FK)", "condition", "status"], "#2980B9"),
        ("DoctorProfile", 9, 1, ["doctor_id (PK/FK)", "specialization", "hospital"], "#2980B9"),
        ("CaregiverProfile", 13, 1, ["caregiver_id (PK/FK)", "alert_channel"], "#2980B9"),
        ("MedicineSchedule", 1, -2.5, ["medicine_id (PK)", "patient_id (FK)", "name", "dosage"], "#27AE60"),
        ("MedicineLog", 1, -5, ["log_id (PK)", "medicine_id (FK)", "status"], "#8E44AD"),
        ("Appointment", 5, -2.5, ["appt_id (PK)", "patient_id (FK)", "doctor_id (FK)", "date", "status"], "#27AE60"),
        ("SOSEvent", 9, -2.5, ["sos_id (PK)", "patient_id (FK)", "status", "location"], "#E74C3C"),
        ("Alert", 13, -2.5, ["alert_id (PK)", "patient_id (FK)", "type", "severity"], "#E67E22"),
        ("ChatSession", 5, -5, ["chat_id (PK)", "user_id (FK)", "language"], "#27AE60"),
        ("Notification", 9, -5, ["notif_id (PK)", "recipient_id (FK)", "type", "channel"], "#27AE60"),
        ("AuditLog", 13, -5, ["audit_id (PK)", "actor_id (FK)", "action", "block_hash"], "#8E44AD"),
        ("ConsentRecord", 9, -7.5, ["consent_id (PK)", "patient_id (FK)", "consent_type"], "#8E44AD"),
        ("HealthReport", 5, -7.5, ["report_id (PK)", "patient_id (FK)", "type"], "#27AE60"),
        ("VoiceSession", 1, -7.5, ["voice_id (PK)", "user_id (FK)", "transcription"], "#27AE60"),
        ("CaregiverMapping", 13, -7.5, ["mapping_id (PK)", "caregiver_id", "patient_id"], "#E67E22"),
    ]

    boxes = {}
    for name, x, y, attrs, color in modules:
        boxes[name] = draw_entity_box(ax, x, y, name, attrs, color, width=3.2, attr_h=0.25)

    # Key relationships
    rels = [
        ("User", "PatientProfile", "1:1", "1", "1"),
        ("User", "DoctorProfile", "1:1", "1", "1"),
        ("User", "CaregiverProfile", "1:1", "1", "1"),
        ("PatientProfile", "MedicineSchedule", "1:M", "1", "M"),
        ("MedicineSchedule", "MedicineLog", "1:M", "1", "M"),
        ("PatientProfile", "Appointment", "1:M", "1", "M"),
        ("DoctorProfile", "Appointment", "1:M", "1", "M"),
        ("PatientProfile", "SOSEvent", "1:M", "1", "M"),
        ("PatientProfile", "Alert", "1:M", "1", "M"),
        ("User", "ChatSession", "1:M", "1", "M"),
        ("User", "Notification", "1:M", "1", "M"),
        ("User", "AuditLog", "1:M", "1", "M"),
        ("PatientProfile", "HealthReport", "1:M", "1", "M"),
        ("User", "VoiceSession", "1:M", "1", "M"),
        ("CaregiverProfile", "CaregiverMapping", "1:M", "1", "M"),
    ]
    for (e1, e2, label, c1, c2) in rels:
        if e1 in boxes and e2 in boxes:
            draw_relationship_line(ax, boxes[e1], boxes[e2], label, c1, c2)

    legend_items = [
        mpatches.Patch(color='#2980B9', label='Core Entity'),
        mpatches.Patch(color='#27AE60', label='Supporting Entity'),
        mpatches.Patch(color='#E67E22', label='Junction/Mapping'),
        mpatches.Patch(color='#8E44AD', label='Log/Audit'),
        mpatches.Patch(color='#E74C3C', label='Emergency'),
    ]
    ax.legend(handles=legend_items, loc='lower right', fontsize=8, framealpha=0.9)

    plt.tight_layout()
    master_path = os.path.join(DIAGRAMS_DIR, "00_master_er_diagram.png")
    plt.savefig(master_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    diagrams["master"] = master_path
    print(f"  ✅ Saved: 00_master_er_diagram.png")

    return diagrams


# ──────────────────────────────────────────────────────────────
# Draw.io XML Generation (editable diagrams)
# ──────────────────────────────────────────────────────────────
def generate_drawio_diagram():
    """Generate an editable draw.io XML file with the complete ER diagram."""
    # Entity definitions: (id, label, x, y, attributes)
    entities = [
        ("user", "User", 50, 50, "user_id (PK)\\nemail\\ndisplay_name\\nrole\\nphone\\nlanguage_preference\\nis_active\\ncreated_at"),
        ("patient", "PatientProfile", 400, 50, "patient_id (PK/FK)\\nage\\ngender\\nblood_group\\ncondition\\nallergies[]\\nchronic_conditions[]\\nassigned_doctor_id (FK)\\nassigned_caregiver_id (FK)\\nadherence\\nstatus"),
        ("doctor", "DoctorProfile", 800, 50, "doctor_id (PK/FK)\\nname\\nspecialization\\nhospital\\nexperience_years\\nrating\\nis_accepting_patients"),
        ("caregiver", "CaregiverProfile", 1200, 50, "caregiver_id (PK/FK)\\nphone\\nrelationship_to_patient\\nalert_preference_channel\\nis_available"),
        ("medicine", "MedicineSchedule", 50, 450, "medicine_id (PK)\\npatient_id (FK)\\nname\\ndosage\\nfrequency\\nschedule_slots[]\\nreminder_times[]\\nis_active"),
        ("medlog", "MedicineLog", 50, 800, "log_id (PK)\\nmedicine_id (FK)\\npatient_id (FK)\\nstatus\\nslot\\ntimestamp"),
        ("appointment", "Appointment", 400, 450, "appointment_id (PK)\\npatient_id (FK)\\ndoctor_id (FK)\\ndate\\ntime\\nreason\\nstatus"),
        ("sos", "SOSEvent", 800, 450, "sos_id (PK)\\npatient_id (FK)\\ntimestamp\\nlocation_lat\\nlocation_lng\\nstatus\\nseverity"),
        ("alert", "Alert", 1200, 450, "alert_id (PK)\\ntype\\nseverity\\npatient_id (FK)\\nsymptom\\nstatus\\nassigned_to (FK)"),
        ("chat", "ChatSession", 400, 800, "chat_id (PK)\\nuser_id (FK)\\nlanguage\\nsession_start\\nmessage_count"),
        ("chatmsg", "ChatMessage", 400, 1100, "message_id (PK)\\nchat_id (FK)\\nrole\\ntext\\nsources[]\\nemergency_flag"),
        ("voice", "VoiceSession", 50, 1100, "voice_id (PK)\\nuser_id (FK)\\nlanguage\\ntranscription\\nai_response"),
        ("notif", "Notification", 800, 800, "notification_id (PK)\\nrecipient_id (FK)\\nchannel\\ntype\\nmessage\\nis_read"),
        ("audit", "AuditLog", 1200, 800, "audit_id (PK)\\nactor_id (FK)\\nrecord_type\\naction\\nblock_hash\\nblock_index"),
        ("consent", "ConsentRecord", 800, 1100, "consent_id (PK)\\npatient_id (FK)\\nconsent_type\\ngranted_to\\nconsent_hash"),
        ("report", "HealthReport", 1200, 1100, "report_id (PK)\\npatient_id (FK)\\ntype\\ndata (JSON)\\ngenerated_at"),
        ("mapping", "CaregiverPatientMapping", 1200, 1400, "mapping_id (PK)\\ncaregiver_id (FK)\\npatient_id (FK)\\nrelationship\\nis_active"),
        ("emergency_contact", "EmergencyContact", 400, 1400, "contact_id (PK)\\npatient_id (FK)\\ncontact_name\\ncontact_phone\\nrelationship"),
        ("clinical_note", "ClinicalNote", 800, 1400, "note_id (PK)\\ndoctor_id (FK)\\npatient_id (FK)\\nnote_content\\ndiagnosis_codes[]"),
    ]

    # Relationships: (source_id, target_id, label, cardinality)
    relationships = [
        ("user", "patient", "1:1"),
        ("user", "doctor", "1:1"),
        ("user", "caregiver", "1:1"),
        ("patient", "medicine", "1:M"),
        ("medicine", "medlog", "1:M"),
        ("patient", "appointment", "1:M"),
        ("doctor", "appointment", "1:M"),
        ("patient", "sos", "1:M"),
        ("patient", "alert", "1:M"),
        ("user", "chat", "1:M"),
        ("chat", "chatmsg", "1:M"),
        ("user", "voice", "1:M"),
        ("user", "notif", "1:M"),
        ("user", "audit", "1:M"),
        ("patient", "consent", "1:M"),
        ("patient", "report", "1:M"),
        ("caregiver", "mapping", "1:M"),
        ("patient", "emergency_contact", "1:M"),
        ("doctor", "clinical_note", "1:M"),
        ("patient", "clinical_note", "1:M"),
    ]

    # Color coding
    color_map = {
        "user": "#2980B9", "patient": "#2980B9", "doctor": "#2980B9", "caregiver": "#2980B9",
        "medicine": "#27AE60", "appointment": "#27AE60", "chat": "#27AE60", "voice": "#27AE60",
        "notif": "#27AE60", "report": "#27AE60", "emergency_contact": "#27AE60", "chatmsg": "#27AE60",
        "medlog": "#8E44AD", "audit": "#8E44AD", "consent": "#8E44AD", "clinical_note": "#27AE60",
        "sos": "#E74C3C", "alert": "#E67E22", "mapping": "#E67E22",
    }

    xml_cells = []
    cell_id = 2

    # Generate entity cells
    for eid, label, x, y, attrs in entities:
        color = color_map.get(eid, "#2980B9")
        w = 280
        h = 30 + len(attrs.split("\\n")) * 18
        xml_cells.append(
            f'      <mxCell id="{cell_id}" value="{label}&#10;──────────────&#10;{attrs}" '
            f'style="shape=table;startSize=30;container=0;collapsible=0;childLayout=tableLayout;'
            f'fillColor={color};fontColor=#FFFFFF;strokeColor=#1A5276;rounded=1;arcSize=10;'
            f'fontSize=11;fontStyle=1;whiteSpace=wrap;html=0;align=left;spacingLeft=8;" '
            f'vertex="1" parent="1">\n'
            f'        <mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry" />\n'
            f'      </mxCell>'
        )
        cell_id += 1

    # Generate relationship edges
    entity_id_map = {eid: i + 2 for i, (eid, *_) in enumerate(entities)}
    for src, tgt, card in relationships:
        src_id = entity_id_map.get(src)
        tgt_id = entity_id_map.get(tgt)
        if src_id and tgt_id:
            xml_cells.append(
                f'      <mxCell id="{cell_id}" value="{card}" '
                f'style="edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;'
                f'jettySize=auto;html=0;strokeColor=#2C3E50;strokeWidth=1.5;'
                f'fontColor=#8E44AD;fontSize=10;fontStyle=2;labelBackgroundColor=#F5EEF8;" '
                f'edge="1" source="{src_id}" target="{tgt_id}" parent="1">\n'
                f'        <mxGeometry relative="1" as="geometry" />\n'
                f'      </mxCell>'
            )
            cell_id += 1

    cells_xml = "\n".join(xml_cells)

    drawio_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<mxfile host="app.diagrams.net" type="device">
  <diagram id="medassist-er" name="MedAssist AI - Complete ER Diagram">
    <mxGraphModel dx="1800" dy="1200" grid="1" gridSize="10" guides="1" tooltips="1"
                  connect="1" arrows="1" fold="1" page="1" pageScale="1"
                  pageWidth="2400" pageHeight="1800" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
{cells_xml}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>
'''
    drawio_path = os.path.join(DIAGRAMS_DIR, "MedAssist_AI_ER_Diagram.drawio")
    with open(drawio_path, "w", encoding="utf-8") as f:
        f.write(drawio_xml)
    print(f"  ✅ Saved: MedAssist_AI_ER_Diagram.drawio (editable)")
    return drawio_path


# ──────────────────────────────────────────────────────────────
# Word Document Generation
# ──────────────────────────────────────────────────────────────
def generate_word_document(diagram_paths):
    """Generate the complete Word (.docx) document."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # ═══════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MedAssist AI")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Intelligent Multilingual Healthcare Assistant")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    doc.add_paragraph()

    doctype = doc.add_paragraph()
    doctype.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doctype.add_run("DATABASE DESIGN DOCUMENT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info table
    info_data = [
        ("Project Title", "MedAssist AI — Intelligent Multilingual Healthcare Assistant"),
        ("Document Type", "Database Design Document"),
        ("Database Platform", "Google Cloud Firestore (NoSQL Document Database)"),
        ("Backend Framework", "FastAPI (Python 3.10+)"),
        ("Authentication", "Firebase Authentication (Email / Password)"),
        ("Version", "2.0.0"),
        ("Prepared By", "MedAssist AI Engineering Team"),
        ("Date", "July 2026"),
    ]
    add_styled_table(doc, ["Field", "Details"], info_data, col_widths=[5, 12])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS PAGE
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. List of Entities",
        "2. List of Attributes",
        "3. Entity-Relationship Diagrams",
        "4. Relationships Between Entities",
        "5. Database Design — Collection Schemas",
        "    5.1. User Management Collections",
        "    5.2. Patient Health Profile Collections",
        "    5.3. Doctor Professional Collections",
        "    5.4. Caregiver Monitoring Collections",
        "    5.5. Medicine Management Collections",
        "    5.6. Appointment Booking Collections",
        "    5.7. Emergency SOS Collections",
        "    5.8. AI Chat and Voice Collections",
        "    5.9. Notification Engine Collections",
        "    5.10. Alert Management Collections",
        "    5.11. Reports and Analytics Collections",
        "    5.12. Audit and Compliance Collections",
        "    5.13. System Configuration Collections",
        "6. Firestore Security Rules",
        "7. Composite Index Requirements",
        "8. Data Integrity and Business Rules",
        "9. Complete Collection Summary",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 1: LIST OF ENTITIES
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. List of Entities", level=1)

    entity_modules = {
        "1. User Management Module": ["User", "UserRole", "UserPreference", "SecurityCredential", "SessionLog", "AccountStatus"],
        "2. Patient Health Profile Module": ["PatientProfile", "EmergencyContact", "Allergy", "ChronicCondition", "BloodGroup", "PatientStatus"],
        "3. Doctor Professional Module": ["DoctorProfile", "Specialization", "HospitalAffiliation", "DoctorAvailability", "ClinicalNote"],
        "4. Caregiver Monitoring Module": ["CaregiverProfile", "CaregiverPatientMapping", "AlertPreference", "CaregiverRelationshipType"],
        "5. Medicine Management Module": ["MedicineSchedule", "MedicineLog", "MedicineSlot", "ReminderConfiguration", "WHOEssentialMedicine", "MedicineCategory"],
        "6. Appointment Booking Module": ["Appointment", "AppointmentStatus", "AppointmentSlot", "AppointmentReason", "AppointmentPolicy"],
        "7. Emergency SOS Module": ["SOSEvent", "SOSNotification", "EmergencyService", "SOSStatus"],
        "8. AI Chat and Voice Module": ["ChatSession", "ChatMessage", "VoiceSession", "RAGQueryLog", "ChatLanguageSetting"],
        "9. Notification Engine Module": ["Notification", "NotificationChannel", "NotificationTemplate", "NotificationPreference", "DeliveryLog"],
        "10. Alert Management Module": ["Alert", "AlertType", "AlertSeverity", "AlertStatus"],
        "11. Reports and Analytics Module": ["HealthReport", "AdherenceReport", "SystemReport", "ReportType", "ReportSchedule"],
        "12. Audit and Compliance Module": ["AuditLog", "ConsentRecord", "BlockchainBlock", "DataAccessLog"],
        "13. System Configuration Module": ["SystemSetting", "LanguageLocale", "APIConfiguration", "RateLimitPolicy", "FeatureFlag"],
    }

    for mod_name, entities in entity_modules.items():
        add_heading_styled(doc, mod_name, level=3)
        for entity in entities:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(entity)
            run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 2: LIST OF ATTRIBUTES (selected key entities)
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. List of Attributes", level=1)

    attribute_sections = {
        "1. User Management Module": {
            "User": [
                ("user_id", "Firebase UID — unique identifier for each user"),
                ("email", "Registered email address of the user"),
                ("display_name", "Full name of the user as displayed in the interface"),
                ("role", "User role: patient / doctor / caregiver / admin"),
                ("phone", "Contact phone number with country code"),
                ("language_preference", "Preferred language code: en / hi / ta / kn / ml / te"),
                ("profile_photo_url", "URL of the user's profile picture"),
                ("is_active", "Whether the user account is currently active"),
                ("email_verified", "Whether the email has been verified via Firebase"),
                ("created_at", "Timestamp of account creation"),
                ("updated_at", "Timestamp of last profile update"),
                ("last_login", "Timestamp of last successful login"),
                ("deleted", "Soft delete flag"),
            ],
        },
        "2. Patient Health Profile Module": {
            "PatientProfile": [
                ("patient_id", "Reference to user_id — one-to-one relationship"),
                ("age", "Patient's age in years"),
                ("gender", "Patient's gender: Male / Female / Other"),
                ("blood_group", "Blood group: A+ / A- / B+ / B- / AB+ / AB- / O+ / O-"),
                ("condition", "Primary medical condition description"),
                ("allergies", "Array of known allergies"),
                ("chronic_conditions", "Array of chronic conditions (diabetes, hypertension)"),
                ("assigned_doctor_id", "Reference to the assigned doctor user_id"),
                ("assigned_caregiver_id", "Reference to the assigned caregiver user_id"),
                ("health_score", "Calculated health score 0–100"),
                ("adherence", "Medicine adherence percentage 0–100"),
                ("status", "Patient status: Active / Stable / Needs Attention / Critical"),
            ],
        },
        "3. Doctor Professional Module": {
            "DoctorProfile": [
                ("doctor_id", "Reference to user_id — one-to-one relationship"),
                ("specialization", "Medical specialization of the doctor"),
                ("hospital", "Name of the affiliated hospital or clinic"),
                ("experience_years", "Number of years of professional experience"),
                ("qualification", "Medical qualifications and degrees"),
                ("license_number", "Medical license or registration number"),
                ("rating", "Average patient rating 0.0–5.0"),
                ("is_accepting_patients", "Whether the doctor is accepting new patients"),
            ],
        },
        "5. Medicine Management Module": {
            "MedicineSchedule": [
                ("medicine_id", "Unique identifier for the medicine schedule entry"),
                ("patient_id", "Reference to the patient this medicine belongs to"),
                ("name", "Name of the medicine"),
                ("dosage", "Dosage amount and unit: 500mg / 10ml"),
                ("frequency", "Administration frequency: Once daily / Twice daily"),
                ("schedule_slots", "Array of time slots: morning / afternoon / evening / night"),
                ("reminder_times", "Array of specific reminder times: 08:00 / 20:00"),
                ("is_active", "Whether the medicine schedule is currently active"),
            ],
        },
        "6. Appointment Booking Module": {
            "Appointment": [
                ("appointment_id", "Unique identifier for the appointment"),
                ("patient_id", "Reference to the patient user_id"),
                ("doctor_id", "Reference to the doctor user_id"),
                ("date", "Appointment date in ISO format YYYY-MM-DD"),
                ("time", "Appointment time e.g. 10:00 AM"),
                ("reason", "Reason for the appointment provided by the patient"),
                ("status", "Scheduled / Confirmed / Completed / Missed / Cancelled"),
            ],
        },
        "7. Emergency SOS Module": {
            "SOSEvent": [
                ("sos_id", "Unique identifier for the SOS event"),
                ("patient_id", "Reference to the patient who triggered the SOS"),
                ("timestamp", "Date and time the SOS was triggered"),
                ("location_latitude", "GPS latitude coordinate if available"),
                ("location_longitude", "GPS longitude coordinate if available"),
                ("status", "SOS status: Active / Acknowledged / Resolved / False Alarm"),
                ("severity", "Severity level: Low / Medium / High / Critical"),
                ("notified_contacts", "Array of user_ids who were notified"),
            ],
        },
        "12. Audit and Compliance Module": {
            "AuditLog": [
                ("audit_id", "Unique identifier for the audit log entry"),
                ("record_type", "Type: medical_record / consent / prescription / access_log"),
                ("actor_id", "User_id of the user who performed the audited action"),
                ("action", "Action: create / read / update / delete / verify"),
                ("data_hash", "SHA-256 hash of the record data for integrity"),
                ("previous_hash", "Hash of the previous audit log — blockchain chaining"),
                ("block_hash", "Combined hash for immutability verification"),
                ("block_index", "Sequential index of this block in the audit chain"),
            ],
        },
    }

    for section_name, entities in attribute_sections.items():
        add_heading_styled(doc, section_name, level=2)
        for entity_name, attrs in entities.items():
            add_heading_styled(doc, entity_name, level=3)
            add_styled_table(doc, ["Attribute", "Description"], attrs, col_widths=[5, 12])
            doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 3: ER DIAGRAMS (embedded images)
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Entity-Relationship Diagrams", level=1)

    p = doc.add_paragraph()
    p.add_run("The following Entity-Relationship (ER) diagrams visually represent the database structure "
              "of MedAssist AI across all 13 modules. Each diagram shows entities (collections), "
              "their key attributes, and the relationships between them including cardinality (1:1, 1:M, M:M).").font.size = Pt(10)
    doc.add_paragraph()

    # Master diagram
    add_heading_styled(doc, "3.0 Complete ER Overview", level=2)
    if "master" in diagram_paths and os.path.exists(diagram_paths["master"]):
        doc.add_picture(diagram_paths["master"], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Module diagrams
    module_diagram_names = [
        ("user_management_er", "3.1 Module 1: User Management"),
        ("patient_profile_er", "3.2 Module 2: Patient Health Profile"),
        ("doctor_er", "3.3 Module 3: Doctor Professional"),
        ("caregiver_er", "3.4 Module 4: Caregiver Monitoring"),
        ("medicine_er", "3.5 Module 5: Medicine Management"),
        ("appointment_er", "3.6 Module 6: Appointment Booking"),
        ("sos_er", "3.7 Module 7: Emergency SOS"),
        ("chat_er", "3.8 Module 8: AI Chat and Voice"),
        ("notification_er", "3.9 Module 9: Notification Engine"),
        ("alert_er", "3.10 Module 10: Alert Management"),
        ("reports_er", "3.11 Module 11: Reports and Analytics"),
        ("audit_er", "3.12 Module 12: Audit and Compliance"),
    ]

    for key, heading in module_diagram_names:
        add_heading_styled(doc, heading, level=2)
        if key in diagram_paths and os.path.exists(diagram_paths[key]):
            doc.add_picture(diagram_paths[key], width=Inches(6.2))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 4: RELATIONSHIPS
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Relationships Between Entities", level=1)

    rel_groups = {
        "User Management": [
            ("UserRole", "User", "1 to M", "One role assigned to many users"),
            ("User", "UserPreference", "1 to 1", "Each user has one preference record"),
            ("User", "SessionLog", "1 to M", "One user has many login sessions"),
            ("AccountStatus", "User", "1 to M", "One status applies to many accounts"),
        ],
        "Patient Health Profile": [
            ("User", "PatientProfile", "1 to 1", "Each patient user has one profile"),
            ("PatientProfile", "EmergencyContact", "1 to M", "One patient has many emergency contacts"),
            ("BloodGroup", "PatientProfile", "1 to M", "One blood group applies to many patients"),
            ("PatientProfile", "DoctorProfile", "M to 1", "Many patients assigned to one doctor"),
        ],
        "Doctor Professional": [
            ("User", "DoctorProfile", "1 to 1", "Each doctor user has one profile"),
            ("Specialization", "DoctorProfile", "1 to M", "One specialization for many doctors"),
            ("DoctorProfile", "ClinicalNote", "1 to M", "One doctor writes many notes"),
        ],
        "Caregiver Monitoring": [
            ("User", "CaregiverProfile", "1 to 1", "Each caregiver has one profile"),
            ("CaregiverProfile", "PatientProfile", "M to M", "Many-to-many via CaregiverPatientMapping"),
        ],
        "Medicine Management": [
            ("PatientProfile", "MedicineSchedule", "1 to M", "One patient has many schedules"),
            ("MedicineSchedule", "MedicineLog", "1 to M", "One schedule generates many logs"),
            ("DoctorProfile", "MedicineSchedule", "1 to M", "One doctor prescribes many schedules"),
        ],
        "Appointment Booking": [
            ("PatientProfile", "Appointment", "1 to M", "One patient books many appointments"),
            ("DoctorProfile", "Appointment", "1 to M", "One doctor has many appointments"),
            ("Appointment", "ClinicalNote", "1 to 1", "Each appointment may have one note"),
        ],
        "Emergency SOS": [
            ("PatientProfile", "SOSEvent", "1 to M", "One patient triggers many SOS events"),
            ("SOSEvent", "Notification", "1 to M", "One SOS sends many notifications"),
        ],
        "AI Chat and Voice": [
            ("User", "ChatSession", "1 to M", "One user has many chat sessions"),
            ("ChatSession", "ChatMessage", "1 to M", "One session has many messages"),
            ("User", "VoiceSession", "1 to M", "One user has many voice sessions"),
            ("ChatMessage", "RAGQueryLog", "1 to 1", "Each response has one RAG log"),
        ],
        "Alert and Notification": [
            ("PatientProfile", "Alert", "1 to M", "One patient has many alerts"),
            ("User", "Notification", "1 to M", "One user receives many notifications"),
            ("SOSEvent", "Alert", "1 to 1", "Each SOS generates one critical alert"),
        ],
        "Audit and Compliance": [
            ("User", "AuditLog", "1 to M", "One user generates many audit entries"),
            ("PatientProfile", "ConsentRecord", "1 to M", "One patient grants many consents"),
            ("AuditLog", "AuditLog", "1 to 1", "Blockchain-style hash chaining"),
        ],
    }

    for group_name, rels in rel_groups.items():
        add_heading_styled(doc, group_name, level=3)
        add_styled_table(doc, ["Entity A", "Entity B", "Cardinality", "Description"],
                        [(r[0], r[1], r[2], r[3]) for r in rels],
                        col_widths=[4, 4, 2.5, 6.5], header_color="8E44AD")
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 5: DATABASE DESIGN — Collection Schemas
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Database Design — Firestore Collection Schemas", level=1)

    p = doc.add_paragraph()
    p.add_run("Firestore organises data into collections (equivalent to tables) and documents "
              "(equivalent to rows). Each document contains fields (equivalent to columns).").font.size = Pt(10)

    # Define all collections
    collections = {
        "5.1 User Management": {
            "users": [
                ("user_id (doc ID)", "String", "PK, Firebase UID", "Unique user identifier"),
                ("email", "String", "Required, Unique, Indexed", "Email address"),
                ("display_name", "String", "Required", "Full name"),
                ("role", "String", "Required, Enum", "patient/doctor/caregiver/admin"),
                ("phone", "String", "Optional", "Phone with country code"),
                ("language_preference", "String", "Default: 'en'", "en/hi/ta/kn/ml/te"),
                ("is_active", "Boolean", "Default: true", "Account active flag"),
                ("created_at", "Timestamp", "Auto-set", "Creation timestamp"),
                ("updated_at", "Timestamp", "Auto-set", "Update timestamp"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
        },
        "5.2 Patient Health Profile": {
            "patients": [
                ("patient_id (doc ID)", "String", "PK, FK → users", "Same as user_id"),
                ("age", "Number", "Range [0-150]", "Patient's age"),
                ("gender", "String", "Enum", "Male/Female/Other"),
                ("blood_group", "String", "Enum", "A+/A-/B+/B-/AB+/AB-/O+/O-"),
                ("condition", "String", "Required", "Primary condition"),
                ("allergies", "Array<String>", "Optional", "Known allergies"),
                ("chronic_conditions", "Array<String>", "Optional", "Chronic conditions"),
                ("assigned_doctor_id", "String", "FK → doctors", "Assigned doctor"),
                ("assigned_caregiver_id", "String", "FK → caregivers", "Assigned caregiver"),
                ("adherence", "Number", "Range [0-100]", "Medicine adherence %"),
                ("status", "String", "Enum", "Active/Stable/Needs Attention/Critical"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
        },
        "5.3 Doctor Professional": {
            "doctors": [
                ("doctor_id (doc ID)", "String", "PK, FK → users", "Same as user_id"),
                ("name", "String", "Required", "Doctor's full name"),
                ("specialization", "String", "Required", "Specialization"),
                ("hospital", "String", "Required", "Affiliated hospital"),
                ("experience_years", "Number", "Range [0-60]", "Years of experience"),
                ("rating", "Number", "Range [0.0-5.0]", "Average rating"),
                ("is_accepting_patients", "Boolean", "Default: true", "Accepting patients"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
            "clinical_notes": [
                ("note_id (doc ID)", "String", "PK, Auto-gen", "Note identifier"),
                ("doctor_id", "String", "FK → doctors, Indexed", "Author doctor"),
                ("patient_id", "String", "FK → patients, Indexed", "Subject patient"),
                ("note_content", "String", "Required", "Clinical note text"),
                ("diagnosis_codes", "Array<String>", "Optional", "ICD-11 codes"),
                ("prescriptions", "Array<Map>", "Optional", "Prescribed medicines"),
            ],
        },
        "5.4 Caregiver Monitoring": {
            "caregivers": [
                ("caregiver_id (doc ID)", "String", "PK, FK → users", "Same as user_id"),
                ("phone", "String", "Optional", "Contact phone"),
                ("alert_preference_channel", "String", "Enum", "push/sms/whatsapp"),
                ("is_available", "Boolean", "Default: true", "Availability"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
            "caregiver_patient_mappings": [
                ("mapping_id (doc ID)", "String", "PK, Auto-gen", "Mapping identifier"),
                ("caregiver_id", "String", "FK → caregivers, Indexed", "Caregiver ref"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient ref"),
                ("relationship", "String", "Required", "Relationship type"),
                ("is_active", "Boolean", "Default: true", "Active flag"),
            ],
        },
        "5.5 Medicine Management": {
            "medicines": [
                ("medicine_id (doc ID)", "String", "PK, Auto-gen", "Medicine identifier"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient reference"),
                ("name", "String", "Required", "Medicine name"),
                ("dosage", "String", "Required", "Dosage: 500mg"),
                ("frequency", "String", "Required", "Frequency description"),
                ("reminder_times", "Array<String>", "Default: []", "Reminder times"),
                ("is_active", "Boolean", "Default: true", "Active flag"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
            "medicine_logs": [
                ("log_id (doc ID)", "String", "PK, Auto-gen", "Log identifier"),
                ("medicine_id", "String", "FK → medicines, Indexed", "Medicine ref"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient ref"),
                ("status", "String", "Enum [taken, missed]", "Dose status"),
                ("slot", "String", "Enum", "morning/afternoon/evening/night"),
                ("timestamp", "Timestamp", "Required", "Event timestamp"),
            ],
        },
        "5.6 Appointment Booking": {
            "appointments": [
                ("appointment_id (doc ID)", "String", "PK, Auto-gen", "Appointment ID"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient reference"),
                ("patient_name", "String", "Required", "Patient display name"),
                ("doctor_id", "String", "FK → doctors, Indexed", "Doctor reference"),
                ("doctor_name", "String", "Required", "Doctor display name"),
                ("date", "String", "Format: YYYY-MM-DD", "Appointment date"),
                ("time", "String", "Required", "Appointment time"),
                ("status", "String", "Enum", "Scheduled/Confirmed/Completed/Cancelled"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
        },
        "5.7 Emergency SOS": {
            "sos_events": [
                ("sos_id (doc ID)", "String", "PK, Auto-gen", "SOS identifier"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient reference"),
                ("timestamp", "Timestamp", "Required", "SOS trigger time"),
                ("location_latitude", "Number", "Optional", "GPS latitude"),
                ("location_longitude", "Number", "Optional", "GPS longitude"),
                ("status", "String", "Enum", "Active/Acknowledged/Resolved"),
                ("severity", "String", "Enum", "Low/Medium/High/Critical"),
                ("notified_contacts", "Array<String>", "Default: []", "Notified IDs"),
            ],
        },
        "5.8 AI Chat and Voice": {
            "chats": [
                ("chat_id (doc ID)", "String", "PK, Auto-gen", "Chat session ID"),
                ("user_id", "String", "FK → users, Indexed", "User reference"),
                ("language", "String", "Enum", "en/hi/ta/kn/ml/te"),
                ("session_start", "Timestamp", "Required", "Session start"),
                ("message_count", "Number", "Default: 0", "Message count"),
                ("model_used", "String", "Default", "AI model identifier"),
            ],
            "voice_sessions": [
                ("voice_id (doc ID)", "String", "PK, Auto-gen", "Voice session ID"),
                ("user_id", "String", "FK → users, Indexed", "User reference"),
                ("language", "String", "Enum", "en/hi/ta/kn/ml/te"),
                ("transcription", "String", "Required", "Whisper transcription"),
                ("ai_response", "String", "Required", "AI response text"),
            ],
        },
        "5.9 Notification Engine": {
            "notifications": [
                ("notification_id (doc ID)", "String", "PK, Auto-gen", "Notification ID"),
                ("recipient_id", "String", "FK → users, Indexed", "Recipient"),
                ("channel", "String", "Enum", "push/sms/whatsapp/ivr/email"),
                ("type", "String", "Enum", "medicine_reminder/appointment/sos/alert"),
                ("message", "String", "Required", "Content text"),
                ("is_read", "Boolean", "Default: false", "Read status"),
            ],
        },
        "5.10 Alert Management": {
            "alerts": [
                ("alert_id (doc ID)", "String", "PK, Auto-gen", "Alert identifier"),
                ("type", "String", "Enum", "Emergency/Missed Medicine/Warning"),
                ("severity", "String", "Enum", "Low/Medium/High/Critical"),
                ("patient_name", "String", "Required", "Patient display name"),
                ("symptom", "String", "Required", "Alert description"),
                ("status", "String", "Enum", "unread/read/Resolved/Dismissed"),
                ("deleted", "Boolean", "Default: false", "Soft delete flag"),
            ],
        },
        "5.11 Reports and Analytics": {
            "reports": [
                ("report_id (doc ID)", "String", "PK, Auto-gen", "Report identifier"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient reference"),
                ("type", "String", "Enum", "Health Summary/Lab Results/Adherence"),
                ("data", "Map (JSON)", "Optional", "Report payload"),
                ("generated_at", "Timestamp", "Optional", "Generation time"),
            ],
        },
        "5.12 Audit and Compliance": {
            "audit_logs": [
                ("audit_id (doc ID)", "String", "PK, Auto-gen", "Audit identifier"),
                ("record_type", "String", "Enum", "medical_record/consent/prescription"),
                ("actor_id", "String", "FK → users, Indexed", "Acting user"),
                ("action", "String", "Enum", "create/read/update/delete/verify"),
                ("data_hash", "String", "Optional", "SHA-256 hash"),
                ("block_hash", "String", "Required", "Block hash"),
                ("block_index", "Number", "Required", "Sequential index"),
                ("verified", "Boolean", "Default: true", "Integrity flag"),
            ],
            "consent_records": [
                ("consent_id (doc ID)", "String", "PK, Auto-gen", "Consent ID"),
                ("patient_id", "String", "FK → patients, Indexed", "Patient ref"),
                ("consent_type", "String", "Enum", "data_sharing/treatment/research"),
                ("granted_to", "String", "Required", "Entity receiving consent"),
                ("consent_hash", "String", "Required", "SHA-256 hash"),
                ("is_active", "Boolean", "Default: true", "Active flag"),
            ],
        },
        "5.13 System Configuration": {
            "system_settings": [
                ("setting_id (doc ID)", "String", "PK, Key name", "Setting identifier"),
                ("setting_key", "String", "Required, Unique", "Setting key"),
                ("setting_value", "Dynamic", "Required", "Setting value"),
                ("category", "String", "Required", "Setting category"),
            ],
            "language_locales": [
                ("locale_id (doc ID)", "String", "PK, Lang code", "Locale ID"),
                ("language_code", "String", "Unique", "ISO code"),
                ("language_name", "String", "Required", "English name"),
                ("native_name", "String", "Required", "Native script name"),
                ("is_active", "Boolean", "Default: true", "Enabled flag"),
            ],
        },
    }

    for section_name, tables in collections.items():
        doc.add_page_break()
        add_heading_styled(doc, section_name, level=2)
        for coll_name, fields in tables.items():
            add_heading_styled(doc, f"Collection: {coll_name}", level=3)
            add_styled_table(doc, ["Field", "Data Type", "Constraints", "Description"],
                           fields, col_widths=[4.5, 2.5, 4, 6], header_color="1B4F72")
            doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 6: SECURITY RULES
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Firestore Security Rules (RBAC)", level=1)

    p = doc.add_paragraph()
    p.add_run("Role-Based Access Control (RBAC) is enforced through Firestore Security Rules. "
              "Only authenticated users can access data, and operations are restricted by user role.").font.size = Pt(10)
    doc.add_paragraph()

    security_rules = [
        ("users", "Owner or Admin", "Owner", "Owner or Admin", "Admin only"),
        ("patients", "Owner / Doctor / Caregiver / Admin", "Patient or Admin", "Owner / Doctor / Admin", "Admin only"),
        ("doctors", "All authenticated", "Doctor or Admin", "Owner or Admin", "Admin only"),
        ("medicines", "All authenticated", "Patient / Doctor / Admin", "Patient / Doctor / Admin", "Patient or Admin"),
        ("appointments", "All authenticated", "Patient or Admin", "Patient / Doctor / Admin", "Admin only"),
        ("sos_events", "All authenticated", "Patient only", "Doctor / Caregiver / Admin", "Not allowed"),
        ("alerts", "All authenticated", "All authenticated", "Doctor / Caregiver / Admin", "Admin only"),
        ("audit_logs", "Admin only", "All authenticated", "Not allowed (immutable)", "Not allowed"),
        ("chats", "All authenticated", "All authenticated", "All authenticated", "Not allowed"),
        ("reports", "All authenticated", "Doctor or Admin", "Doctor or Admin", "Admin only"),
        ("system_settings", "All authenticated", "Admin only", "Admin only", "Admin only"),
    ]

    add_styled_table(doc, ["Collection", "Read", "Create", "Update", "Delete"],
                    security_rules, col_widths=[3, 3.5, 3.5, 3.5, 3], header_color="C0392B")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 7: INDEX REQUIREMENTS
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Composite Index Requirements", level=1)

    indexes = [
        ("patients", "assigned_doctor_id ASC, status ASC", "Doctor's patients by status"),
        ("patients", "assigned_caregiver_id ASC, status ASC", "Caregiver's patients by status"),
        ("medicines", "patient_id ASC, is_active ASC", "Active medicines for a patient"),
        ("medicine_logs", "patient_id ASC, timestamp DESC", "Recent logs for a patient"),
        ("appointments", "patient_id ASC, date DESC", "Patient's appointments by date"),
        ("appointments", "doctor_id ASC, date ASC, status ASC", "Doctor's appointments"),
        ("sos_events", "patient_id ASC, timestamp DESC", "SOS history for a patient"),
        ("sos_events", "status ASC, timestamp DESC", "Active SOS events"),
        ("chats", "user_id ASC, session_start DESC", "User's chat sessions"),
        ("alerts", "patient_id ASC, status ASC, created_at DESC", "Patient alerts"),
        ("notifications", "recipient_id ASC, is_read ASC, created_at DESC", "Unread notifications"),
        ("audit_logs", "actor_id ASC, timestamp DESC", "Audit trail by user"),
        ("audit_logs", "record_type ASC, timestamp DESC", "Audit trail by type"),
    ]

    add_styled_table(doc, ["Collection", "Index Fields", "Query Purpose"],
                    indexes, col_widths=[3.5, 6, 7.5], header_color="27AE60")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 8: BUSINESS RULES
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Data Integrity and Business Rules", level=1)

    rules_sections = {
        "Referential Integrity Rules": [
            "Every patient_id MUST reference a valid user with role = 'patient'",
            "Every doctor_id MUST reference a valid user with role = 'doctor'",
            "Every caregiver_id MUST reference a valid user with role = 'caregiver'",
            "Appointment patient_id and doctor_id MUST reference existing documents",
        ],
        "Soft Delete Rules": [
            "All primary collections implement soft delete via 'deleted' Boolean field",
            "Soft-deleted documents are excluded from all list queries",
            "Hard delete is restricted to admin users (GDPR compliance)",
        ],
        "Timestamp Rules": [
            "created_at is auto-set on creation and MUST NOT be modified",
            "updated_at is auto-set on every document update",
            "All timestamps use ISO 8601 format in UTC timezone",
        ],
        "Status Transition Rules": [
            "Appointment: Scheduled → Confirmed → Completed (terminal) OR Cancelled (terminal)",
            "SOS: Active → Acknowledged → Resolved (terminal) OR False Alarm (terminal)",
            "Alert: unread → read → Resolved (terminal) OR Dismissed (terminal)",
            "Terminal statuses MUST NOT transition to any other status",
        ],
        "Data Validation Rules": [
            "Email addresses MUST conform to RFC 5322 format",
            "Phone numbers MUST include country code prefix",
            "Age MUST be in range 0–150",
            "Adherence percentage MUST be in range 0–100",
            "Health score MUST be in range 0–100",
            "Doctor rating MUST be in range 0.0–5.0",
        ],
    }

    for section_name, rules in rules_sections.items():
        add_heading_styled(doc, section_name, level=3)
        for rule in rules:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(rule).font.size = Pt(10)
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # SECTION 9: COLLECTION SUMMARY TABLE
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Complete Collection Summary", level=1)

    summary = [
        ("1", "users", "User Management", "Thousands"),
        ("2", "user_roles", "User Management", "4 (static)"),
        ("3", "user_preferences", "User Management", "Thousands"),
        ("4", "session_logs", "User Management", "Tens of thousands"),
        ("5", "patients", "Patient Profile", "Hundreds–Thousands"),
        ("6", "doctors", "Doctor Profile", "Tens–Hundreds"),
        ("7", "caregivers", "Caregiver Profile", "Tens–Hundreds"),
        ("8", "caregiver_patient_mappings", "Caregiver Monitoring", "Hundreds"),
        ("9", "clinical_notes", "Doctor Module", "Thousands"),
        ("10", "medicines", "Medicine Management", "Thousands"),
        ("11", "medicine_logs", "Medicine Management", "Tens of thousands"),
        ("12", "who_essential_medicines", "Medicine Management", "1,738 (WHO EML)"),
        ("13", "appointments", "Appointment Booking", "Thousands"),
        ("14", "appointment_slots", "Appointment Booking", "Thousands"),
        ("15", "sos_events", "Emergency SOS", "Hundreds"),
        ("16", "chats", "AI Chat", "Thousands"),
        ("17", "voice_sessions", "AI Voice", "Thousands"),
        ("18", "rag_query_logs", "AI Pipeline", "Tens of thousands"),
        ("19", "notifications", "Notifications", "Tens of thousands"),
        ("20", "alerts", "Alerts", "Thousands"),
        ("21", "reports", "Reports", "Thousands"),
        ("22", "adherence_reports", "Reports", "Thousands"),
        ("23", "audit_logs", "Audit", "Tens of thousands"),
        ("24", "consent_records", "Audit", "Hundreds"),
        ("25", "data_access_logs", "Audit", "Tens of thousands"),
        ("26", "system_settings", "System Config", "Tens (static)"),
        ("27", "language_locales", "System Config", "6 (static)"),
        ("28", "api_configurations", "System Config", "6–10 (static)"),
    ]

    add_styled_table(doc, ["#", "Collection Name", "Module", "Est. Documents"],
                    summary, col_widths=[1, 5.5, 4, 4], header_color="1B4F72")

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("━" * 60)
    run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("End of Database Design Document — MedAssist AI v2.0.0")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.italic = True

    # ── Save ──
    doc.save(DOCX_PATH)
    print(f"\n✅ Word document saved: {DOCX_PATH}")


# ──────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("MedAssist AI — Database Design Document Generator")
    print("=" * 60)

    print("\n📊 Generating ER Diagrams...")
    diagram_paths = generate_all_diagrams()

    print("\n📐 Generating Editable Draw.io Diagram...")
    generate_drawio_diagram()

    print("\n📝 Generating Word Document (.docx)...")
    generate_word_document(diagram_paths)

    print("\n" + "=" * 60)
    print("✅ All files generated successfully!")
    print(f"   📄 Word Document: {DOCX_PATH}")
    print(f"   📊 ER Diagrams:   {DIAGRAMS_DIR}/")
    print(f"   📐 Editable:      {DIAGRAMS_DIR}/MedAssist_AI_ER_Diagram.drawio")
    print("=" * 60)
